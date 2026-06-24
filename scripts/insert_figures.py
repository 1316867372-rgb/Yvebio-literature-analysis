# -*- coding: utf-8 -*-
"""将Fig图片和流程图嵌入8模块docx对应位置"""

import docx, os, sys
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def insert_images(docx_path, image_dir):
    doc = Document(docx_path)
    docx_name = os.path.splitext(os.path.basename(docx_path))[0]
    prefix = (
        docx_name.split("_文献拆解_")[0] if "_文献拆解_" in docx_name else docx_name
    )

    positions = {}
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if "五、研究流程" in text:
            positions["flow"] = i + 1
        for fig_num in range(1, 10):
            if text.startswith(f"Fig {fig_num}") or text.startswith(f"Fig{fig_num}"):
                positions[f"fig{fig_num}"] = i

    items = []
    all_files = os.listdir(image_dir)
    # flow image
    flow_candidates = [
        f for f in all_files if prefix in f and "研究流程" in f and f.endswith(".png")
    ]
    if "flow" in positions and flow_candidates:
        items.append(
            ("flow", positions["flow"], os.path.join(image_dir, flow_candidates[0]))
        )
    # fig images
    for fig_num in range(9, 0, -1):
        key = f"fig{fig_num}"
        if key in positions:
            candidates = [
                f for f in all_files if prefix in f and f.endswith(f"Fig.{fig_num}.png")
            ]
            if candidates:
                items.append(
                    (key, positions[key], os.path.join(image_dir, candidates[0]))
                )

    items.sort(key=lambda x: x[1], reverse=True)

    for key, idx, img_path in items:
        para = doc.paragraphs[idx]
        new_para = docx.oxml.shared.OxmlElement("w:p")
        para._element.addnext(new_para)
        new_p = docx.text.paragraph.Paragraph(new_para, doc)
        new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        new_p.add_run().add_picture(img_path, width=Inches(5.5))
        print(f"  [{key}] inserted after para {idx}")

    doc.save(docx_path)
    print(f"Done. {len(items)} images inserted.")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python insert_figures.py <docx_path>")
        sys.exit(1)
    insert_images(sys.argv[1], os.path.dirname(sys.argv[1]))
